from docx import Document
import json

doc = Document(r'D:\先挣一个亿\上传素材\游戏内容介绍说明_先挣一个亿.docx')

# 输出文档结构
structure = []
for para in doc.paragraphs:
    if para.text.strip():
        structure.append({
            'text': para.text,
            'style': para.style.name,
            'heading_level': para.style.name
        })

# 输出表格
tables = []
for i, table in enumerate(doc.tables):
    table_data = []
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        table_data.append(row_data)
    tables.append(table_data)

print("=== 文档结构 ===")
for item in structure:
    print(f"[{item['style']}] {item['text'][:100]}...")
    
print("\n=== 表格 ===")
for i, table in enumerate(tables):
    print(f"\n表格 {i+1}:")
    for row in table:
        print(f"  {row}")