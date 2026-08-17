from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# 创建文档
doc = Document()

# 设置默认字体
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
doc.styles['Normal'].font.size = Pt(11)

# 使用内置标题样式，只修改字体
for i in range(1, 4):
    style_name = f'Heading {i}'
    if style_name in doc.styles:
        doc.styles[style_name].font.name = '微软雅黑'
        doc.styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 主标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('《赛季全胜》')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(24)
run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('游戏内容介绍说明')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(16)

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('（开发者版）')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# 目录
p = doc.add_paragraph('目录', style='Heading 1')

toc_items = [
    '一、游戏总体介绍',
    '二、游戏背景',
    '三、主要角色',
    '四、场景介绍',
    '五、主要情节',
    '六、玩法',
    '七、功能（系统）介绍',
    '八、道具、装备说明',
    '九、技能说明',
    '十、概念神系统'
]

for item in toc_items:
    p = doc.add_paragraph(item)

doc.add_page_break()

# ============ 一、游戏总体介绍 ============
p = doc.add_paragraph('一、游戏总体介绍', style='Heading 1')

p = doc.add_paragraph('1. 游戏总体介绍（必要项）', style='Heading 2')

p = doc.add_paragraph('《赛季全胜》是一款以"招募球员积累财富"为核心的竖屏放置挂机类微信小游戏。在游戏中，玩家通过不断消耗预算进行招募，获得不同品质的NBA风格球员卡片，利用积累的球员阵容提高球队总评，挑战越来越强的对手赛季。')

p = doc.add_paragraph('游戏遵循"纯增量、无惩罚"的设计哲学——比赛失败不会扣除任何资源，也不会影响正式战绩，玩家可以通过招募强化阵容后从失败场次重新挑战，最终建立常规赛82–0、季后赛16–0的无敌赛季。')

p = doc.add_paragraph('玩家可以通过观看激励视频广告获得各种增益效果：免费招募、开启自动翻转、强制胜利、离线收益翻倍、直接获得预算、管理层免费升级等。')

p = doc.add_paragraph('游戏类型：放置挂机（Idle Clicker）+ 模拟经营  目标平台：微信小游戏（IAA）  屏幕方向：竖屏  目标用户：篮球爱好者、上班摸鱼/学生党、男性为主，对NBA有基础认知的玩家')

# 插入logo图片
try:
    logo_path = r'D:\篮球\上传素材\赛季全胜游戏logo设计_(1).png'
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(4))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# ============ 二、游戏背景 ============
p = doc.add_paragraph('二、游戏背景', style='Heading 1')

p = doc.add_paragraph('2. 游戏背景（次要项）', style='Heading 2')

p = doc.add_paragraph('你是一支NBA球队的老板，手握球队的全部预算和管理权。你的目标是通过精明的球员招募策略，建立一支无敌的冠军阵容。')

p = doc.add_paragraph('从新秀赛季开始，你将经历13个不同难度的品质赛季，从最初的新秀阵容开始，逐步招募更高品质的球员，最终挑战GOAT级别的传奇赛季。')

p = doc.add_paragraph('每一次赛季夺冠后，你可以提升球队市值，解锁更高品质的球员招募池。与此同时，下个赛季的对手也会变得更加强大。你的终极目标是集齐12名概念神球员，达成游戏毕业，在无限赛程中永远保持不败。')

# ============ 三、主要角色 ============
p = doc.add_paragraph('三、主要角色', style='Heading 1')

p = doc.add_paragraph('3. 主要角色（次要项）', style='Heading 2')

p = doc.add_paragraph('游戏中的核心角色是球员和管理层两大类：')

p = doc.add_paragraph('球员系统：')
p = doc.add_paragraph('游戏包含13档球员品质，从低到高依次为：新秀（绿）、饮水机（绿）、轮换（蓝）、第六人（蓝）、首发（紫）、核心（紫）、全明星（橙）、最佳阵容（橙）、MVP（红）、FMVP（红）、名人堂（粉）、传奇（金）、GOAT（虹彩）。')

p = doc.add_paragraph('每档品质都有独立的OVR区间，品质越高，球员的基础能力越强。球员池包含真实NBA球员原型和CBA大陆MVP/FMVP球星，共计1300张球员卡片。所有球员使用中文同音化名命名，如"酷里"（库里）、"瞻姆斯"（詹姆斯）、"步莱恩特"（布莱恩特）等。')

p = doc.add_paragraph('管理层系统：')
p = doc.add_paragraph('共有5个管理层职位，各自负责不同的专项加成：')
p = doc.add_paragraph('• 运营总裁：提升在线预算收益、比赛预算奖励、广告预算奖励')
p = doc.add_paragraph('• 主教练：提供比赛OVR百分比加成，计入球队总评')
p = doc.add_paragraph('• 球探总监：提升当前球队市值下最高已解锁品质的招募概率权重')
p = doc.add_paragraph('• 队医团队：新招募球员卡OVR随机结果向品质上限正向偏移')
p = doc.add_paragraph('• 媒体团队：提升离线预算收益倍率')

# ============ 四、场景介绍 ============
p = doc.add_paragraph('四、场景介绍', style='Heading 1')

p = doc.add_paragraph('4. 场景介绍（次要项）', style='Heading 2')

p = doc.add_paragraph('游戏整体采用现代轻奢金色主题设计，界面风格简洁大气。主要包含以下界面：')

# 主界面
p = doc.add_paragraph('主游戏界面', style='Heading 3')
p = doc.add_paragraph('主界面顶部显示当前预算、可领取的离线收益按钮；中部为招募按钮，点击即可消耗预算招募新球员；底部为导航栏，包含招募/赛季/球员/管理层/设置入口。')

# 插入主页截图
try:
    img_path = r'D:\篮球\上传素材\主页.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(3.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# 赛季/备战界面
p = doc.add_paragraph('赛季备战界面', style='Heading 3')
p = doc.add_paragraph('备战界面显示当前赛季进度、常规赛胜场数、下一场对手信息和对手OVR对比。玩家可以在这里查看本场比赛信息、调整阵容，并点击"开始比赛"进入比赛模拟。')

# 插入备战截图
try:
    img_path = r'D:\篮球\上传素材\备战.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(3.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# 比赛界面
p = doc.add_paragraph('比赛界面', style='Heading 3')
p = doc.add_paragraph('比赛界面模拟真实篮球比赛的四节过程，每节30秒。界面显示双方球队名、实时比分、当前节数、比赛文字播报。支持二倍速播放和跳过按钮，玩家可以选择立即查看预生成的赛果。')

# 插入比赛截图
try:
    img_path = r'D:\篮球\上传素材\比赛.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(3.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# 胜利界面
p = doc.add_paragraph('胜利界面', style='Heading 3')
p = doc.add_paragraph('比赛胜利后显示胜利弹窗，展示本场比赛的最终比分、获得的预算奖励。玩家可以选择观看广告使奖励翻倍，或直接继续下一场比赛。')

# 插入胜利截图
try:
    img_path = r'D:\篮球\上传素材\胜利.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(3.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# 招募结果界面
p = doc.add_paragraph('招募结果界面', style='Heading 3')
p = doc.add_paragraph('每次招募后弹出球员卡片，展示新球员的头像、姓名、品质、OVR、五项属性（得分/篮板/助攻/抢断/盖帽）。界面自动对比新球员与当前阵容中OVR最低的球员，突出展示替换后的球队总评绝对增量。玩家可以选择上阵（替换旧球员）或解雇（放弃新球员），也可以观看广告将新球员品质提升一级。')

# 插入招募结果截图
try:
    img_path = r'D:\篮球\上传素材\招募结果.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(3.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# 招募概率界面
p = doc.add_paragraph('招募概率界面', style='Heading 3')
p = doc.add_paragraph('点击招募界面的概率图标可查看当前可招募的各品质概率。界面按品质从低到高排列，展示当前可招募的5档品质及实际抽取百分比。球探区域显示当前球探等级和球探加成带来的最高品质概率增量。')

# 插入招募概率截图
try:
    img_path = r'D:\篮球\上传素材\招募概率.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(3.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# 管理层界面
p = doc.add_paragraph('管理层界面', style='Heading 3')
p = doc.add_paragraph('管理层界面展示5个管理层职位的当前等级、加成效果和升级成本。玩家可以消耗预算升级管理层，或通过观看广告免费升级。每个管理层等级不得超过当前球队等级。')

# 插入管理层截图
try:
    img_path = r'D:\篮球\上传素材\管理层.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(3.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# 球员详情界面
p = doc.add_paragraph('球员详情界面', style='Heading 3')
p = doc.add_paragraph('点击上阵球员可以查看球员详情，展示球员头像、姓名、品质、OVR、五项属性数值和属性占比雷达图。')

# 插入球员详情截图
try:
    img_path = r'D:\篮球\上传素材\球员详情.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(3.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# 离线收益界面
p = doc.add_paragraph('离线收益界面', style='Heading 3')
p = doc.add_paragraph('玩家重新进入游戏时弹出离线收益弹窗，展示离线时长和累计获得的预算。玩家可以直接领取，或观看广告使离线收益翻倍。')

# 插入离线收益截图
try:
    img_path = r'D:\篮球\上传素材\离线收益.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(3.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except:
    pass

# ============ 五、主要情节 ============
p = doc.add_paragraph('五、主要情节', style='Heading 1')

p = doc.add_paragraph('5. 主要情节（次要项）', style='Heading 2')

p = doc.add_paragraph('游戏以赛季为核心循环单位，每个赛季包含常规赛82场和季后赛四轮，玩家需要达成98场全胜才能夺冠并进入下一品质难度的赛季。')

p = doc.add_paragraph('赛季难度从低到高依次为：新秀、饮水机、轮换、第六人、首发、核心、全明星、最佳阵容、MVP、FMVP、名人堂、传奇、GOAT，共13个品质赛季。')

p = doc.add_paragraph('完成GOAT难度赛季98–0并夺冠后，不再开启新的标准赛季，而是永久进入"概念神无限赛程"，比赛编号无限累加，难度持续提升，玩家可以在此模式下挑战游戏的极限。')

# ============ 六、玩法 ============
p = doc.add_paragraph('六、玩法', style='Heading 1')

p = doc.add_paragraph('6. 玩法（必要项）', style='Heading 2')

p = doc.add_paragraph('游戏的核心玩法为"招募球员积累财富→提升球队总评→挑战更高难度赛季"，玩家通过以下循环实现增长：')

p = doc.add_paragraph('核心玩法流程：')
p = doc.add_paragraph('• 点击招募 → 消耗预算获得新球员卡 + 固定获得斗志')
p = doc.add_paragraph('• 新球员卡与当前阵容中OVR最低的球员比较，展示球队总评增量')
p = doc.add_paragraph('• 新球员更强：上阵替换，旧球员自动离队')
p = doc.add_paragraph('• 新球员较弱：放弃新球员卡，不影响现有阵容')
p = doc.add_paragraph('• 招募固定获得斗志 → 斗志满足后点击升级按钮提升球队等级')
p = doc.add_paragraph('• 达到当前球队市值等级上限后，进入季后赛并夺冠 → 球队市值提升')
p = doc.add_paragraph('• 球队市值提升 → 解锁更高品质招募概率 + 提高球队等级上限')
p = doc.add_paragraph('• 赛季全胜夺冠 → 进入下一品质难度赛季')

p = doc.add_paragraph('比赛胜负规则：')
p = doc.add_paragraph('• 若当前12名上阵球员全部为概念神 → 直接获胜')
p = doc.add_paragraph('• 玩家OVR ≥ 对手OVR × 1.1 → 必胜')
p = doc.add_paragraph('• 玩家OVR ≤ 对手OVR × 0.9 → 必败')
p = doc.add_paragraph('• 对手OVR × 0.9 < 玩家OVR < 对手OVR × 1.1 → 50%胜率')

p = doc.add_paragraph('失败后处理：')
p = doc.add_paragraph('• 失败不计入正式战绩、不给比赛奖励、不推进赛程，自动暂停在当前比赛')
p = doc.add_paragraph('• 可选择观看广告获得1%–20%临时球队总评加成并重赛')
p = doc.add_paragraph('• 也可以返回继续招募强化阵容后再次挑战')

# ============ 七、功能（系统）介绍 ============
p = doc.add_paragraph('七、功能（系统）介绍', style='Heading 1')

p = doc.add_paragraph('7. 功能（系统）介绍（必要项）', style='Heading 2')

p = doc.add_paragraph('本游戏包含以下核心系统：')

# 招募系统
p = doc.add_paragraph('（1）招募系统', style='Heading 3')
p = doc.add_paragraph('招募是游戏的核心"开箱"行为。玩家消耗固定预算进行招募，每次招募必定获得一张球员卡，并固定获得10点斗志。')

p = doc.add_paragraph('招募概率由当前球队市值决定，每一级市值对应一张完整的品质招募概率表。球探总监会对当前可招募池内的最高品质追加权重。')

p = doc.add_paragraph('招募结果只与当前12人阵容中OVR最低的球员比较，重点展示替换后的球队总评绝对增量。')

# 阵容系统
p = doc.add_paragraph('（2）阵容系统', style='Heading 3')
p = doc.add_paragraph('玩家拥有12个球员槽位，所有槽位地位等同，不强制要求任何位置。球员自身有位置（PG/SG/SF/PF/C），但位置只作为球员标签和生成五项属性时的初始权重，比赛引擎不读取位置字段。')

p = doc.add_paragraph('球队总评计算公式：')
p = doc.add_paragraph('球员基础总评 = 12名上阵球员OVR之和')
p = doc.add_paragraph('球队总评 = floor(球员基础总评 × (1 + 主教练比赛OVR加成))')

p = doc.add_paragraph('若12名上阵球员全部为概念神：')
p = doc.add_paragraph('球队总评 = 2,147,483,647（显示MAX），所有比赛直接获胜')

# 球队等级与市值系统
p = doc.add_paragraph('（3）球队等级与市值系统', style='Heading 3')
p = doc.add_paragraph('每次招募固定获得10点斗志，斗志达到当前等级需求后，由玩家点击升级按钮提升球队等级。每次点击只提升1级，多余斗志保留。')

p = doc.add_paragraph('球队市值共130级，每级球队市值包含4个球队等级小段。达到当前球队市值阶段的球队等级上限后，即使斗志满足也不能继续升级，需要进入季后赛并夺冠后才能提升市值并解锁下一组球队等级。')

p = doc.add_paragraph('球队市值决定：当前可招募的品质概率表、管理层可升级上限、球队等级上限。')

# 赛季系统
p = doc.add_paragraph('（4）赛季系统', style='Heading 3')
p = doc.add_paragraph('每个赛季包含常规赛82场和季后赛四轮，需要达成98场全胜才能夺冠。赛季难度按已完成冠军赛季数递增，共13个品质赛季。')

p = doc.add_paragraph('常规赛：82场逐场进行，对手OVR线性递增，只有获胜才推进场次，最终必须82–0才能进入季后赛。')

p = doc.add_paragraph('季后赛：四轮对手强度分别为总决赛门槛的85%/90%/95%/100%，每轮需要4胜通关，最终系列赛比分固定为4–0。')

p = doc.add_paragraph('概念神无限赛程：完成GOAT赛季后解锁，取消赛季概念，比赛编号与难度无限递增。')

# 管理层系统
p = doc.add_paragraph('（5）管理层系统', style='Heading 3')

# 创建管理层表格
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'

# 表头
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '管理层'
hdr_cells[1].text = '职能'
hdr_cells[2].text = '每级效果'
hdr_cells[3].text = '520级效果'

# 数据
data = [
    ['运营总裁', '预算运营', '预算收益+0.2%', '+104%'],
    ['主教练', '比赛战术', '比赛OVR+0.05%', '+26%'],
    ['球探总监', '招募品质', '最高品质权重+0.01', '+5.20'],
    ['队医团队', '球员养成', '新卡OVR随机+0.05百分点', '+26百分点'],
    ['媒体团队', '知名度收益', '离线预算+0.4%', '+208%']
]

for i, row_data in enumerate(data):
    row_cells = table.rows[i+1].cells
    for j, cell_text in enumerate(row_data):
        row_cells[j].text = cell_text

p = doc.add_paragraph('')
p = doc.add_paragraph('管理层初始均为0级，最高520级。每个管理层的当前等级不得超过当前球队等级。升级消耗预算，或可通过观看广告免费升级。')

# IAA广告系统
p = doc.add_paragraph('（6）IAA广告系统', style='Heading 3')
p = doc.add_paragraph('游戏接入激励视频广告，共8个广告点位：')

p = doc.add_paragraph('• 免费招募：不消耗预算，完整执行一次正常招募')
p = doc.add_paragraph('• 招募结果升级：普通品质升一级；符合资格的GOAT升级为概念神')
p = doc.add_paragraph('• 强制胜利：0.9–1.1不确定区间随机判负后可改判为胜利')
p = doc.add_paragraph('• 离线收益翻倍：存在可领取离线预算时可使收益×2')
p = doc.add_paragraph('• 直接获得预算：基础奖励=10×当前球队等级')
p = doc.add_paragraph('• 管理层免费升级：任一管理层职位免费提升1级')
p = doc.add_paragraph('• 胜利奖励翻倍：胜利弹窗可追加一份等额奖励')
p = doc.add_paragraph('• 失败强化重赛：失败后获得1%–20%临时总评加成并重赛')

# ============ 八、道具、装备说明 ============
p = doc.add_paragraph('八、道具、装备说明', style='Heading 1')

p = doc.add_paragraph('8. 道具、装备说明（次要项）', style='Heading 2')

p = doc.add_paragraph('本游戏无独立道具装备系统。球员卡片本身既是核心养成对象，也是直接影响球队战力的"装备"。')

p = doc.add_paragraph('管理层职位相当于"被动装备"，提供持续的经济或战力加成。五个管理层职位各管一摊，玩家可以根据自己的发展策略选择优先升级方向。')

# ============ 九、技能说明 ============
p = doc.add_paragraph('九、技能说明', style='Heading 1')

p = doc.add_paragraph('9. 技能说明（次要项）', style='Heading 2')

p = doc.add_paragraph('本游戏无独立技能系统。球员的能力通过五项属性（得分、篮板、助攻、抢断、盖帽）体现，其中得分属性直接影响比赛中的比分表现。')

p = doc.add_paragraph('主教练的战术加成作为"被动技能"，直接影响球队总评；其他管理层的加成效果也相当于各自领域的专项技能。')

# ============ 十、概念神系统 ============
p = doc.add_paragraph('十、概念神系统', style='Heading 1')

p = doc.add_paragraph('10. 概念神系统（高级系统）', style='Heading 2')

p = doc.add_paragraph('概念神是独立于13档普通品质的特殊品质，代表游戏中的终极追求。概念神不进入普通招募权重表，只能由具有概念神资格的GOAT球员在招募结果页完整观看激励视频后升级获得。')

p = doc.add_paragraph('当前共48个概念神，每个都有独特的梗解释和原型来源。概念神在升级成功时生成独立的OVR区间，第X名概念神使用GOAT区间×(1+0.01×X)的倍率。')

p = doc.add_paragraph('首次凑齐12名概念神时永久记录"概念神全阵容·游戏毕业"成就，此后所有比赛直接获胜，球队总评显示为MAX。')

p = doc.add_paragraph('概念神或无概念神资格的GOAT可以继续通过观看广告随机强化一项尚未满值的属性，每项属性最高可达到2,147,483,647。')

# 保存文档
output_path = r'D:\篮球\游戏内容介绍说明_赛季全胜.docx'
doc.save(output_path)
print(f"文档已生成：{output_path}")
